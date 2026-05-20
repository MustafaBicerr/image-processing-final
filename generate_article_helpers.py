"""
IEEE Article Generator - Part 1: Document Setup and Helper Functions
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json


REPO_DIR = os.path.dirname(os.path.abspath(__file__))
RESULTS_DIR = os.path.join(REPO_DIR, 'results')
FIGURES_DIR = os.path.join(RESULTS_DIR, 'figures')
OUTPUT_PATH = os.path.join(REPO_DIR, '..', 'Mustafa_Bicer_Final_Report.docx')


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_ieee_heading(doc, text, level=1):
    """Add IEEE-style heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(10) if level == 1 else Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)
    return heading


def add_body_text(doc, text, bold=False, italic=False):
    """Add body text in IEEE style."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    return p


def add_figure(doc, path, caption, width=Inches(3.2)):
    """Add a figure with caption."""
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.font.size = Pt(8)
        cap_run.font.name = 'Times New Roman'
        cap.paragraph_format.space_after = Pt(6)


def create_results_table(doc, data):
    """Create results comparison table."""
    methods = ['CLAHE-RGB', 'CLAHE-HSI', 'HE-RGB', 'HE-HSI', 'Gamma-RGB', 'Gamma-HSI', 'ACSS']
    headers = ['Method', 'PSNR (dB)', 'SSIM', 'CIEDE2000', 'Runtime (ms)']

    table = doc.add_table(rows=len(methods)+1, cols=len(headers))
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.name = 'Times New Roman'

    for i, method in enumerate(methods):
        if method in data:
            d = data[method]
            vals = [
                method,
                f"{d['PSNR_mean']:.2f} ± {d['PSNR_std']:.2f}",
                f"{d['SSIM_mean']:.4f} ± {d['SSIM_std']:.4f}",
                f"{d['CIEDE2000_mean']:.2f} ± {d['CIEDE2000_std']:.2f}",
                f"{d['Runtime_mean_ms']:.1f}",
            ]
        else:
            vals = [method, '-', '-', '-', '-']

        for j, v in enumerate(vals):
            cell = table.rows[i+1].cells[j]
            cell.text = v
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
                    r.font.name = 'Times New Roman'

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run('TABLE I: Quantitative comparison of enhancement methods on Kodak dataset')
    r.font.size = Pt(8)
    r.font.name = 'Times New Roman'
    r.bold = True
    return table


if __name__ == '__main__':
    print("Helper module loaded. Run generate_article_main.py instead.")
